"""景区知识库 - 基于 ChromaDB 向量存储（支持 txt / pdf / docx / doc / xlsx / xls）"""
import os
import zipfile
import xml.etree.ElementTree as ET

from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.embeddings import DashScopeEmbeddings
from langchain_community.vectorstores import Chroma


def _read_txt(path):
    for enc in ("utf-8", "gbk", "gb18030", "utf-16"):
        try:
            with open(path, "r", encoding=enc) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def _read_docx(path):
    try:
        import docx as _python_docx
        doc = _python_docx.Document(path)
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    except ImportError:
        pass
    try:
        with zipfile.ZipFile(path, "r") as z:
            with z.open("word/document.xml") as f:
                tree = ET.parse(f)
        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        paras = []
        for p in tree.findall(".//w:p", ns):
            texts = [t.text for t in p.findall(".//w:t", ns) if t.text]
            if texts:
                paras.append("".join(texts).strip())
        return "\n".join(paras)
    except Exception as e:
        print(f"  [KB] docx zip fallback fail for {os.path.basename(path)}: {e}")
        return ""


def _read_xlsx(path):
    try:
        import openpyxl
        wb = openpyxl.load_workbook(path, data_only=True, read_only=True)
        out = []
        for ws in wb.worksheets:
            out.append(f"【工作表: {ws.title}】")
            for row in ws.iter_rows(values_only=True):
                cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
                if cells:
                    out.append(" | ".join(cells))
            out.append("")
        return "\n".join(out).strip()
    except Exception as e:
        print(f"  [KB] xlsx read fail for {os.path.basename(path)}: {e}")
        return ""


def _read_xls(path):
    try:
        import xlrd
        wb = xlrd.open_workbook(path)
        out = []
        for ws in wb.sheets():
            out.append(f"【工作表: {ws.name}】")
            for r in range(ws.nrows):
                cells = [str(ws.cell_value(r, c)).strip() for c in range(ws.ncols)
                         if ws.cell_value(r, c) not in (None, "")]
                if cells:
                    out.append(" | ".join(cells))
            out.append("")
        return "\n".join(out).strip()
    except Exception as e:
        print(f"  [KB] xls read fail for {os.path.basename(path)}: {e}")
        return ""


def _read_pdf(path):
    try:
        import pdfplumber
        pages = []
        with pdfplumber.open(path) as pdf:
            for p in pdf.pages:
                t = p.extract_text() or ""
                if t.strip():
                    pages.append(t)
        return "\n".join(pages)
    except Exception as e:
        print(f"  [KB] pdf read fail for {os.path.basename(path)}: {e}")
        return ""


def _read_doc_binary(path):
    """旧版 .doc：二进制，Python 侧无稳健解析器；
    用 olefile + 正则提取 UTF-16LE 段做兜底。"""
    try:
        data = open(path, "rb").read()
        # 取 wordDocument 流里的 UTF-16LE 文本
        idx = data.find(b"\x00\x00\x00\x00")
        chunk = data[min(0, idx): idx if idx > 0 else len(data)]
        import re
        raw = chunk.split(b"\x00\x00")
        out_chunks = []
        for block in raw:
            if len(block) >= 6:
                try:
                    s = block.decode("utf-16le", errors="ignore")
                except Exception:
                    continue
                s = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]+", "", s)
                s = s.replace("\r", " ").strip()
                if s and len(s) > 2:
                    out_chunks.append(s)
        text = "\n".join(out_chunks)
        # 再做一次纯 ascii/utf-8 段的兼容扫描
        ascii_parts = re.findall(rb"[\x20-\x7E\u4e00-\u9fff]{4,}", data)
        ascii_text = "\n".join(p.decode("utf-8", errors="ignore") for p in ascii_parts)
        combined = (text + "\n" + ascii_text).strip()
        if combined:
            return combined
    except Exception as e:
        print(f"  [KB] doc binary fallback fail for {os.path.basename(path)}: {e}")
    return ""


READERS = {
    ".txt": _read_txt,
    ".md": _read_txt,
    ".pdf": _read_pdf,
    ".docx": _read_docx,
    ".doc": _read_doc_binary,
    ".xlsx": _read_xlsx,
    ".xls": _read_xls,
}


class KnowledgeBase:
    ALLOWED_EXTS = set(READERS.keys())

    def __init__(self, data_dir="../data/raw", vector_db_path="../data/chroma_storage"):
        self.data_dir = os.path.abspath(data_dir)
        self.vector_db_path = os.path.abspath(vector_db_path)
        os.makedirs(self.data_dir, exist_ok=True)

        self.embeddings = DashScopeEmbeddings(
            model="text-embedding-v4",
            dashscope_api_key=os.getenv("DASHSCOPE_API_KEY"),
        )

        self.splitter = RecursiveCharacterTextSplitter(
            chunk_size=500,
            chunk_overlap=60,
            separators=["\n\n", "\n", "。", "！", "？", "；", "，", " ", ""],
        )

        self.collection_name = "lingshan_knowledge"
        self.vector_store = self._load_vector_store()

    # ---------- 向量库 ----------
    def _load_vector_store(self):
        if os.path.exists(self.vector_db_path):
            try:
                return Chroma(
                    collection_name=self.collection_name,
                    embedding_function=self.embeddings,
                    persist_directory=self.vector_db_path,
                )
            except Exception as e:
                print(f"  [KB] load chroma failed: {e}")
        return None

    def _reset_vector_store(self):
        """彻底重建：用 chromadb API 删 collection → 重新 from_texts"""
        import chromadb
        try:
            client = chromadb.PersistentClient(path=self.vector_db_path)
            if self.collection_name in [c.name for c in client.list_collections()]:
                client.delete_collection(self.collection_name)
                print(f"  [KB] chromadb deleted collection '{self.collection_name}'")
            else:
                print(f"  [KB] chromadb collection '{self.collection_name}' not present")
        except Exception as e:
            print(f"  [KB] chromadb reset via API failed ({e}), fallback to shutil")
            try:
                from shutil import rmtree
                if os.path.exists(self.vector_db_path):
                    rmtree(self.vector_db_path, ignore_errors=True)
            except Exception as e2:
                print(f"  [KB] shutil fallback also failed: {e2}")
        return None

    # ---------- 文档 ----------
    def list_files(self):
        if not os.path.isdir(self.data_dir):
            return []
        out = []
        for name in sorted(os.listdir(self.data_dir)):
            p = os.path.join(self.data_dir, name)
            if not os.path.isfile(p):
                continue
            ext = os.path.splitext(name)[1].lower()
            if ext not in self.ALLOWED_EXTS:
                continue
            out.append({
                "name": name,
                "ext": ext.lstrip("."),
                "size": os.path.getsize(p),
            })
        return out

    def load_documents(self):
        docs = []
        if not os.path.isdir(self.data_dir):
            print(f"  [KB] data_dir not found: {self.data_dir}")
            return docs

        for name in sorted(os.listdir(self.data_dir)):
            p = os.path.join(self.data_dir, name)
            if not os.path.isfile(p):
                continue
            ext = os.path.splitext(name)[1].lower()
            reader = READERS.get(ext)
            if reader is None:
                continue
            try:
                content = reader(p)
            except Exception as e:
                print(f"  [KB] skip {name}: {e}")
                continue
            if not content or not content.strip():
                print(f"  🚫 {name}: 解析后内容为空")
                continue
            print(f"  📄 {name}: {len(content)} 字 (parser={ext})")
            docs.append({
                "content": content,
                "source": name,
                "ext": ext.lstrip("."),
                "chars": len(content),
            })
        print(f"✅ 共加载 {len(docs)} 个文档")
        return docs

    def build_vector_store(self):
        docs = self.load_documents()

        all_texts, all_metadatas = [], []
        for doc in docs:
            try:
                chunks = self.splitter.split_text(doc["content"])
            except Exception as e:
                print(f"  [KB] split {doc['source']} failed: {e}")
                chunks = [doc["content"]]
            for i, chunk in enumerate(chunks):
                all_texts.append(chunk)
                all_metadatas.append({
                    "source": doc["source"],
                    "chunk_id": i,
                    "ext": doc.get("ext", ""),
                    "doc_chars": doc.get("chars", 0),
                })

        print(f"✂️  切分为 {len(all_texts)} 个文本块")

        # 必须先清空旧库，否则会 append 导致旧残留
        self.vector_store = self._reset_vector_store()

        vector_store = Chroma.from_texts(
            texts=all_texts,
            embedding=self.embeddings,
            metadatas=all_metadatas,
            collection_name=self.collection_name,
            persist_directory=self.vector_db_path,
        )
        self.vector_store = vector_store
        print(f"✅ 向量库构建完成：{len(all_texts)} 条记录")
        return {"chunks": len(all_texts), "docs": len(docs)}

    # ---------- 搜索 ----------
    def search(self, query: str, k: int = 5):
        if self.vector_store is None:
            print("  [KB] 向量库未加载，先重建一次")
            self.build_vector_store()
            if self.vector_store is None:
                return []

        try:
            docs = self.vector_store.similarity_search(query, k=k)
        except Exception as e:
            print(f"  [KB] similarity_search failed: {e}")
            return []

        results = []
        for doc in docs:
            results.append({
                "content": doc.page_content,
                "source": (doc.metadata or {}).get("source", ""),
                "chunk_id": (doc.metadata or {}).get("chunk_id", 0),
            })
        return results
